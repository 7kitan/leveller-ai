"""
gap_v3 nodes: CV Parsing Pipeline (Pipeline 1).
5 nodes: extract_text → llm_parse → normalize → pii_mask → persist.
"""

from ast import Return
import uuid
import os
import logging
from typing import Dict, Any

from ..states import CVParsingState, CVParsedData
from ..utils.pii_masker import mask_pii, mask_work_history
from ..utils.llm_helpers import llm_json_completion
from ..utils.ocr_client import ocr_client
from shared.models import UserCV, UserSkillProfile, Skill, UserWorkExperience
import json
from shared.redis_client import result_cache
from shared.config_utils import config_manager

logger = logging.getLogger("cv_parsing_v3")

def _update_cv_progress(cv_id: str, step_message: str, percent: int):
    """Ghi nhận tiến độ hiện tại vào Redis để Frontend có thể polling."""
    try:
        data = json.dumps({"step": step_message, "percent": percent})
        result_cache.set(f"cv_progress:{cv_id}", data, ex=3600)
    except Exception as e:
        logger.warning(f"Failed to update progress in Redis: {e}")


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from .docx file using python-docx library.
    
    Args:
        file_path: Path to .docx file
        
    Returns:
        Extracted text content
        
    Raises:
        Exception: If extraction fails
    """
    try:
        from docx import Document
        
        logger.info(f"[DOCX] Extracting text from: {file_path}")
        doc = Document(file_path)
        
        # Extract text from all paragraphs
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text.strip())
        
        extracted_text = "\n".join(text_parts)
        logger.info(f"[DOCX] Extracted {len(extracted_text)} characters from {len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables")
        
        return extracted_text
        
    except ImportError:
        logger.error("[DOCX] python-docx library not installed. Run: pip install python-docx")
        raise Exception("python-docx library not available")
    except Exception as e:
        logger.error(f"[DOCX] Failed to extract text: {e}")
        raise


# ─── Node 1: Text Extraction ─────────────────────────────────────────────


async def extract_text_node(state: CVParsingState) -> CVParsingState:
    """
    STEP 1: Extract raw text từ CV file.
    Strategy:
      1. Nếu DB đã có raw_text (> 100 chars) → dùng cache, skip extraction
      2. Dùng pymupdf đọc text thuần từ PDF
      3. Nếu text < 200 chars → OCR fallback via pdf2image
    """
    cv_id_str = state["cv_id"]
    db = state["db"]

    logger.info(
        f"\n{'─' * 50}\n"
        f"[STEP 1] extract_text_node | cv_id={cv_id_str}\n"
        f"  DB session : {db}"
    )
    _update_cv_progress(cv_id_str, "Đang trích xuất nội dung văn bản (OCR/PDF)...", 10)

    # ── Validate cv_id ───────────────────────────────────────────────────────
    try:
        cv_uuid = uuid.UUID(cv_id_str)
        logger.info(f"[STEP 1] cv_id valid UUID: {cv_uuid}")
    except ValueError as e:
        logger.error(f"[STEP 1] INVALID cv_id={cv_id_str}: {e}")
        return {**state, "error": f"Invalid cv_id: {cv_id_str}", "status": "failed"}

    # ── Load CV record ───────────────────────────────────────────────────────
    cv_record = db.query(UserCV).filter(UserCV.id == cv_uuid).first()
    if not cv_record:
        logger.error(f"[STEP 1] CV record NOT FOUND in DB: cv_id={cv_id_str}")
        return {**state, "error": f"CV not found: {cv_id_str}", "status": "failed"}
    logger.info(
        f"[STEP 1] CV record found | file_id={cv_record.file_id} | "
        f"file_hash={cv_record.file_hash[:8] if cv_record.file_hash else 'N/A'}... | "
        f"existing raw_text len={len(cv_record.raw_text or '')}"
    )

    # ── Locate CV file on disk ───────────────────────────────────────────────
    file_id = getattr(cv_record, "file_id", None) or cv_id_str
    upload_dir = os.getenv("CV_UPLOAD_DIR", "data/cv_uploads")
    
    # Strategy Detection
    strategy = config_manager.get_setting("CV_PARSER_STRATEGY", default="direct").lower()
    logger.info(f"[STEP 1] Selected Strategy: {strategy.upper()}")

    # Robust File Discovery
    file_path = None
    if os.path.exists(upload_dir):
        for f_name in os.listdir(upload_dir):
            if f_name.startswith(f"{file_id}."):
                file_path = os.path.join(upload_dir, f_name)
                logger.info(f"[STEP 1] ✓ File DISCOVERED on disk: {file_path}")
                break
    
    # BUG-001 FIX: Cache hit validation - Check both raw_text AND cv_parsed_json exist
    # If raw_text exists but cv_parsed_json is missing, re-parse is needed
    if cv_record.raw_text and len(cv_record.raw_text) > 100:
        # Additional validation: check if cv_parsed_json exists
        if cv_record.cv_parsed_json:
            logger.info(
                f"[STEP 1] ✓ CACHE HIT — using existing raw_text "
                f"({len(cv_record.raw_text)} chars) | is_ocr={getattr(cv_record, 'is_ocr', False)}"
            )
            return {
                **state,
                "raw_text": cv_record.raw_text,
                "is_ocr": getattr(cv_record, "is_ocr", False),
                "status": "text_extracted",
                "file_path": file_path,
            }
        else:
            logger.warning(
                f"[STEP 1] CACHE PARTIAL — raw_text exists but cv_parsed_json missing. "
                f"Will re-extract and re-parse."
            )
    
    logger.info("[STEP 1] CACHE MISS — extracting text from file...")
    
    try:
        import fitz  # pymupdf
        
        if not file_path:
            logger.error(f"[STEP 1] FILE NOT FOUND on disk starting with: {file_id}")
            return {
                **state,
                "error": f"CV file not found for ID: {file_id} in {upload_dir}",
                "status": "failed",
            }

        file_size = os.path.getsize(file_path)
        logger.info(f"[STEP 1] File ready | size={file_size:,} bytes | path={file_path}")

    except Exception as e:
        logger.error(f"[STEP 1] Error locating file: {e}", exc_info=True)
        return {**state, "error": f"Error locating CV file: {e}", "status": "failed"}

    # ── Execution: Chandra vs Direct ──────────────────────────────────────
    is_ocr = False
    
    # Check file extension to determine extraction method
    file_ext = os.path.splitext(file_path)[1].lower()
    
    # Handle .docx/.doc files
    if file_ext in ['.docx', '.doc']:
        logger.info(f"[STEP 1] DOCX/DOC file detected: {file_ext}")
        try:
            raw_text = extract_text_from_docx(file_path)
            
            if len(raw_text) < 100:
                logger.error(f"[STEP 1] DOCX extraction failed: text too short ({len(raw_text)} chars)")
                return {
                    **state,
                    "error": "Validation failed: The document does not contain sufficient CV elements.",
                    "status": "failed"
                }
            
            logger.info(f"[STEP 1] ✓ DOCX SUCCESS | text_len={len(raw_text)}")
            
            # Save to DB for caching
            cv_record.raw_text = raw_text
            cv_record.is_ocr = False
            db.commit()
            
            return {
                **state,
                "raw_text": raw_text,
                "is_ocr": False,
                "status": "text_extracted",
                "file_path": file_path,
            }
        except Exception as e:
            logger.error(f"[STEP 1] DOCX extraction failed: {e}", exc_info=True)
            return {
                **state,
                "error": f"Failed to extract text from DOCX: {str(e)}",
                "status": "failed"
            }

    if strategy == "chandra":
        logger.info(f"[STEP 1] HUB REQUEST: Sending {os.path.basename(file_path)} to Chandra OCR API...")
        ocr_result = await ocr_client.ocr_file(file_path)
        
        if ocr_result["status"] == "success":
            raw_text = ocr_result["text"]
            is_ocr = ocr_result["is_ocr"]
            logger.info(
                f"[STEP 1] ✓ CHANDRA SUCCESS | text_len={len(raw_text)} | "
                f"is_ocr={is_ocr}"
            )
            return {
                **state,
                "raw_text": raw_text,
                "is_ocr": is_ocr,
                "status": "text_extracted",
                "file_path": file_path,
            }
        else:
            logger.warning(
                f"[STEP 1] ⚠ CHANDRA HUB FAILED: {ocr_result.get('error')} | "
                f"Falling back to DIRECT extraction..."
            )
            # Continue to direct logic

    # ── Extract text via pymupdf (DIRECT) ────────────────────────────────────
    raw_text = ""
    file_ext = os.path.splitext(file_path)[1].lower()
    is_pdf = file_ext in (".pdf", "pdf")

    if is_pdf:
        try:
            doc = fitz.open(file_path)
            num_pages = len(doc)
            logger.info(f"[STEP 1] Opening PDF with pymupdf | pages={num_pages}")

            text_pages = []
            for page_num in range(num_pages):
                page = doc[page_num]
                text = page.get_text("text")
                text_len = len(text.strip()) if text else 0
                logger.info(
                    f"  page {page_num + 1}/{num_pages}: {text_len} chars extracted"
                )
                if text and text_len > 50:
                    text_pages.append(text)

            doc.close()
            raw_text = "\n\n".join(text_pages)

            logger.info(
                f"[STEP 1] Text extraction complete:\n"
                f"  pages with text  : {len(text_pages)}/{num_pages}\n"
                f"  total raw_text   : {len(raw_text)} chars\n"
                f"  first 120 chars  : {repr(raw_text[:120])}"
            )
        except Exception as e:
            logger.warning(f"[STEP 1] pymupdf extraction failed (will try OCR fallback): {e}")
            raw_text = ""
    else:
        logger.info(f"[STEP 1] File is not PDF ({file_ext}), skipping pymupdf, going to OCR fallback.")

    # BUG-002 FIX: OCR fallback threshold - Lowered from 200 to 100 chars
    # Some valid single-page CVs may have 150-180 chars, so 200 was too high
    is_ocr = False
    if len(raw_text.strip()) < 100:
        logger.warning(
            f"[STEP 1] ⚠ LOW TEXT YIELD ({len(raw_text)} chars < 100) — "
            f"attempting OCR fallback..."
        )
        try:
            from PIL import Image
            images = []
            
            if is_pdf:
                import pdf2image
                dpi = config_manager.get_setting("OCR_DPI", default=200, cast=int)
                logger.info(f"[STEP 1] pdf2image converting PDF at dpi={dpi} ...")
                images = pdf2image.convert_from_path(file_path, dpi=dpi)
                logger.info(f"[STEP 1] Converted {len(images)} images from PDF pages")
            else:
                logger.info(f"[STEP 1] Opening image file directly for OCR: {file_path}")
                images = [Image.open(file_path)]

            ocr_parts = []
            for idx, img in enumerate(images, 1):
                ocr_parts.append(f"[OCR PAGE {idx}: size={img.size}]")
                logger.info(f"  page {idx}: {img.size[0]}x{img.size[1]} px")

            raw_text = "\n\n".join(ocr_parts) + "\n\n" + raw_text
            is_ocr = True
            logger.info(
                f"[STEP 1] OCR metadata markers applied | is_ocr=True | "
                f"final text length={len(raw_text)} chars"
            )

        except ImportError:
            logger.warning("[STEP 1] pdf2image not installed — skipping OCR fallback")
        except Exception as ocr_err:
            logger.warning(
                f"[STEP 1] OCR fallback failed (continuing with partial text): {ocr_err}"
            )

    # ── Validate result ───────────────────────────────────────────────────────
    if not raw_text.strip():
        logger.error(f"[STEP 1] ✗ FAIL — CV file is empty or unreadable: {file_path}")
        return {
            **state,
            "error": "CV file is empty or unreadable. Please check your file and try again.",
            "status": "failed",
        }

    logger.info(
        f"[STEP 1] ✓ SUCCESS | is_ocr={is_ocr} | "
        f"extracted={len(raw_text)} chars | cv_id={cv_id_str}\n"
        f"{'─' * 50}"
    )
    return {
        **state,
        "raw_text": raw_text,
        "is_ocr": is_ocr,
        "status": "text_extracted",
    }


# ─── Node 2: LLM Structured Parsing ────────────────────────────────────


async def llm_parse_cv_node(state: CVParsingState) -> CVParsingState:
    """
    STEP 2: LLM Structured Parsing.
    Dùng OpenAI GPT parse raw_text → structured CVParsedData.
    PII được mask trước khi gửi LLM.
    """
    cv_id_str = state["cv_id"]
    raw_text = state.get("raw_text", "")
    is_ocr = state.get("is_ocr", False)

    logger.info(
        f"\n{'─' * 50}\n"
        f"[STEP 2] llm_parse_cv_node | cv_id={cv_id_str}\n"
        f"  raw_text length : {len(raw_text)} chars\n"
        f"  is_ocr          : {is_ocr}\n"
        f"  text preview    : {repr(raw_text[:200])}"
    )
    _update_cv_progress(cv_id_str, "Hệ thống AI đang phân tích cấu trúc CV (Bước này mất 1-3 phút)...", 30)

    from datetime import datetime
    current_date = datetime.now().strftime("%Y-%m-%d")

    if not raw_text or len(raw_text) < 50:
        logger.error(
            f"[STEP 2] ✗ FAIL — raw_text too short ({len(raw_text)} chars < 50)"
        )
        return {**state, "error": "Raw text too short or empty", "status": "failed"}

    # BUG-003 FIX: PII Masking is now MANDATORY for GDPR compliance
    # Removed ENABLE_PII_MASKING flag - always mask PII before sending to LLM
    logger.info("[STEP 2] Masking PII before sending to LLM (MANDATORY for GDPR compliance)...")
    masked_text = mask_pii(raw_text)

    logger.info(
        f"[STEP 2] PII masking complete | before={len(raw_text)} chars | "
        f"after={len(masked_text)} chars"
    )

    # ── Build prompt ──────────────────────────────────────────────────────────
    prompt = f"""
    SYSTEM ROLE:
    You are a Precision HR Data Architect. Your task is to:
    1. Validate if the uploaded text is a Curriculum Vitae (CV) or Resume.
    2. If it is a CV, transform it into a high-fidelity JSON.
    3. If it is NOT a CV, return a specific failure status.

    TODAY'S DATE: {current_date}

    VALIDATION RULE:
    - A document is considered a CV if it contains at least TWO of the following: Full Name, Contact Info, Education History, Work Experience, or Professional Skills.
    - If the document is an invoice, a random article, a book chapter, or any non-CV text, set "status": "fail" and stop extraction.

    STRICT RULES (Only apply if document is a CV):
    1. FACTUAL INTEGRITY: Extract ONLY information explicitly present. Do not infer skills.
    2. DATE PRECISION & OVERLAP LOGIC: 
       - Use {current_date} for any "Present", "Now", or "Current" end dates.
       - NON-ADDITIVE CALCULATION: Identify all unique time segments for total experience.
       - For individual jobs ('duration_years'), calculate the exact decimal years using ONLY the dates associated directly with that specific job. Example: Jan 2023 to May 2026 is (2026-2023) + (5-1)/12 = 3.3 years. 
       - DO NOT hallucinate. Do NOT use dates from subsequent lines (e.g. do not mix project dates with work history dates).
    3. LANGUAGE: All summaries and descriptions must be translated into English.
    4. NO NORMALIZATION: Keep 'raw_name' for technical skills (e.g., "Py" remains "Py").
    5. CONTEXTUAL SENIORITY: 
       - Evaluated seniority based on RELEVANT experience to the target role.
       - Junior (< 2 yrs), Mid-level (2-5 yrs), Senior (5-10 yrs), Expert (> 10 yrs).
    6. MESSY TEXT PROTOCOL: Use "Visual Block Anchor" to link dates to job titles within the same logical section.
    7. SKILL EXPERIENCE CALCULATION LOGIC:
       - For each skill identified, scan the 'Work History' and 'Projects' sections.
       - If a skill (or its synonym) is mentioned in a job description or project description, attribute the duration of that job/project to the skill's 'experience_years'.
       - If a skill is only listed in a standalone 'Skills' section without a timeframe, set its 'experience_years' equal to the duration of the most recent relevant professional role.
       - Apply the same NON-ADDITIVE CALCULATION (Rule 2) to skills to ensure overlapping roles don't double-count years for a single skill.
    8. CERTIFICATIONS & LICENSES:
       - Extract any mentioned certificates, professional licenses, or language proficiencies (e.g., DELF, IELTS, AWS Certified) into the 'certifications' list.
    9. OCR SPACING RECONSTRUCTION:
       - The OCR text may have spaces between every single letter and number (e.g., 'J a n  2 0 2 3  -  N o w  P o w e r'). 
       - You MUST carefully reconstruct these characters into words ('Jan 2023 - Now Power'). Do not skip these spaced-out dates. Use them as the official start/end dates for the adjacent job title.
    10. AUTO-GENERATED SUMMARY:
       - If the CV does not explicitly contain a summary or objective, DO NOT return null for the 'summary' field. You MUST auto-generate a concise professional summary in English based on the candidate's work history and skills.
    11. SKILL LEVEL EVALUATION:
       - For each skill, deduce its 'level' based on its calculated 'experience_years' using the scale: Junior (< 2 yrs), Mid-level (2-5 yrs), Senior (5-10 yrs), Expert (> 10 yrs).

    INTERNAL MONOLOGUE:
    - Step 0: [Validation] Does this text look like a CV? If no, prepare "fail" response.
    - Step 1: Chronological Audit (List dates, subtract overlaps).
    - Step 2: Relevance Filter for Seniority.
    - Step 3: Skill-to-Role Mapping.
    - Step 4: Quality Check for 'ocr_confidence'.
    - Step 5: Skill Duration Trace (Map every skill to specific time blocks in work history/projects to calculate experience_years).

    ## CV TEXT:
    {masked_text}

    ## OUTPUT SCHEMA (DO NOT CHANGE ANY KEYS):
    {{
      "status": "success | fail",
      "error_message": "Reason if fail, else null",
      "full_name": "Full Name or null",
      "summary": "Professional summary in English (auto-generated if missing)",
      "seniority": "Junior | Mid-level | Senior | Expert | null",
      "experience_years_total": 0.0,
      "skills": [
        {{
          "name": "Skill Name",
          "category": "Technology | Tool | Programming Language | Library | AI | Framework | etc.",
          "experience_years": 0.0,
          "level": "Junior | Mid-level | Senior | Expert"
        }}
      ],
      "work_history": [
        {{
          "position": "Title",
          "company": "Company Name",
          "duration_years": 0.0,
          "description": "Short description in English"
        }}
      ],
      "education": [
        {{
          "degree": "...",
          "institution": "...",
          "year": 2024
        }}
      ],
      "certifications": ["Cert A", "Cert B"],
      "ocr_confidence": 0.0
    }}

    IMPORTANT: Return ONLY valid JSON.
    """

    # ── LLM call ───────────────────────────────────────────────────────────────
    from ..config import GAP_LLM_MODEL as LLM_MODEL, ENABLE_PII_MASKING

    logger.info(
        f"[STEP 2] Calling LLM: model={LLM_MODEL} | "
        f"input chars={len(prompt)} | is_ocr={is_ocr}"
    )
    t_llm = __import__("time").monotonic()
    user_id = state.get("user_id")
    result = await llm_json_completion(
        prompt, 
        context=f"is_ocr={is_ocr}", 
        model_key="cv_parsing_model",
        user_id=user_id
    )
    elapsed_llm = __import__("time").monotonic() - t_llm
    logger.info(
        f"[STEP 2] LLM returned in {elapsed_llm:.1f}s | result keys={list(result.keys()) if result else 'EMPTY'}"
    )

    if not result:
        logger.error(
            f"[STEP 2] ✗ FAIL — LLM returned empty result after {elapsed_llm:.1f}s\n"
            f"  is_ocr={is_ocr} | raw_text chars={len(raw_text)}"
        )
        return {
            **state,
            "error": f"LLM parsing returned empty result (took {elapsed_llm:.1f}s)",
            "status": "failed",
        }

    # ── Validation: Check if the document is actually a CV ────────────────────
    if result.get("status") == "fail":
        error_msg = result.get("error_message") or "Document is not a valid CV/Resume"
        logger.warning(f"[STEP 2] ✗ VALIDATION FAILED — {error_msg}")
        return {
            **state,
            "error": f"Validation failed: {error_msg}",
            "status": "failed",
        }

    # ── Build CVParsedData ────────────────────────────────────────────────────
    parsed: CVParsedData = {
        "full_name": result.get("full_name") or "Không xác định",
        "summary": result.get("summary") or "",
        "seniority": result.get("seniority") or "Unknown",
        "experience_years_total": float(result.get("experience_years_total") or 0),
        "skills": result.get("skills") or [],
        "work_history": result.get("work_history") or [],
        "education": result.get("education") or [],
        "certifications": result.get("certifications") or [],
        "is_ocr": is_ocr,
        "ocr_confidence": float(
            result.get("ocr_confidence") or (0.6 if is_ocr else 1.0)
        ),
        # Removed raw_text_masked as per user request to keep JSON clean
    }

    skill_count = len(parsed.get("skills", []))
    work_count = len(parsed.get("work_history", []))
    edu_count = len(parsed.get("education", []))
    cert_count = len(parsed.get("certifications", []))

    logger.info(
        f"[STEP 2] ✓ SUCCESS\n"
        f"  full_name       : {parsed['full_name']}\n"
        f"  seniority       : {parsed['seniority']}\n"
        f"  experience_yrs  : {parsed['experience_years_total']}\n"
        f"  skills          : {skill_count} entries\n"
        f"  work_history    : {work_count} entries\n"
        f"  education       : {edu_count} entries\n"
        f"  certifications  : {cert_count} entries\n"
        f"  is_ocr          : {is_ocr}\n"
        f"  ocr_confidence  : {parsed['ocr_confidence']:.2f}\n"
        f"  LLM elapsed     : {elapsed_llm:.1f}s\n"
        f"{'─' * 50}"
    )

    return {**state, "cv_parsed": parsed, "status": "parsed"}


# ─── Node 3: Normalize + Validate ───────────────────────────────────────


async def normalize_cv_node(state: CVParsingState) -> CVParsingState:
    """
    STEP 3: Normalize + Validate parsed CV data.
    - Deduplicate skill names
    - Normalize skill levels
    - Normalize seniority labels
    - Clamp experience years to non-negative
    """
    cv_id_str = state["cv_id"]
    cv_parsed = state.get("cv_parsed")

    logger.info(
        f"\n{'─' * 50}\n"
        f"[STEP 3] normalize_cv_node | cv_id={cv_id_str}\n"
        f"  cv_parsed keys : {list(cv_parsed.keys()) if cv_parsed else 'None'}"
    )
    _update_cv_progress(cv_id_str, "Đang chuẩn hóa và sàng lọc bộ kỹ năng...", 80)

    if not cv_parsed:
        logger.error(f"[STEP 3] ✗ FAIL — cv_parsed is None/empty")
        return {**state, "status": "failed", "error": "No parsed CV data"}

    # ── Normalize Skills ──────────────────────────────────────────────────────
    level_map = {
        "beginner": "Junior",
        "intern": "Junior",
        "intermediate": "Mid-level",
        "mid": "Mid-level",
        "advanced": "Senior",
        "expert": "Expert",
        "senior": "Senior",
        "junior": "Junior",
        "lead": "Expert",
    }
    seniority_map = {
        "junior": "Junior",
        "mid": "Mid-level",
        "mid-level": "Mid-level",
        "middle": "Mid-level",
        "senior": "Senior",
        "expert": "Expert",
        "lead": "Expert",
    }

    normalized_skills = []
    seen_names: set = set()
    skill_dedup_count = 0

    for idx, skill in enumerate(cv_parsed.get("skills", [])):
        raw_name = (skill.get("name") or "").strip()
        if not raw_name:
            logger.debug(f"  skill[{idx}] skipped: empty name")
            continue

        # BUG-004 FIX: Normalize skill name to title case for consistency
        # This ensures "python", "Python", "PYTHON" all become "Python"
        name_normalized = raw_name.title() if len(raw_name) > 2 else raw_name.upper()
        name_lower = name_normalized.lower()
        
        if name_lower in seen_names:
            skill_dedup_count += 1
            logger.debug(f"  skill[{idx}] deduplicated: '{raw_name}' (normalized: '{name_normalized}')")
            continue

        raw_level = (skill.get("level") or "Junior").lower().strip()
        normalized = level_map.get(raw_level, "Junior")

        # Normalize category to English
        raw_cat = (skill.get("category") or "Technology").strip()
        cat_lower = raw_cat.lower()
        if cat_lower == "công nghệ":
            raw_cat = "Technology"

        normalized_skills.append(
            {
                "name": name_normalized,  # BUG-004 FIX: Use normalized name
                "category": raw_cat,
                "experience_years": max(0.0, float(skill.get("experience_years") or skill.get("years_exp") or 0)),
                "level": normalized,
            }
        )
        seen_names.add(name_lower)
        logger.debug(
            f"  skill[{idx}] kept: name='{raw_name}' → '{name_normalized}' | "
            f"level='{skill.get('level')}' → '{normalized}' | "
            f"yrs_exp={skill.get('experience_years', 0)}"
        )

    cv_parsed["skills"] = normalized_skills

    # ── Normalize Seniority ───────────────────────────────────────────────────
    raw_senior = (cv_parsed.get("seniority") or "Unknown").lower().strip()
    final_senior = seniority_map.get(raw_senior, raw_senior.title())
    cv_parsed["seniority"] = final_senior
    logger.info(f"[STEP 3] Seniority: '{raw_senior}' → '{final_senior}'")

    # ── Clamp Experience Years ────────────────────────────────────────────────
    raw_exp = float(cv_parsed.get("experience_years_total") or 0)
    cv_parsed["experience_years_total"] = max(0.0, raw_exp)

    # ── Log Summary ──────────────────────────────────────────────────────────
    logger.info(
        f"[STEP 3] ✓ COMPLETE\n"
        f"  skills before dedup : {len(state['cv_parsed'].get('skills', []))}\n"
        f"  skills deduplicated : {skill_dedup_count}\n"
        f"  skills after norm   : {len(normalized_skills)}\n"
        f"  seniority           : {final_senior}\n"
        f"  experience_yrs      : {cv_parsed['experience_years_total']}\n"
        f"{'─' * 50}"
    )

    return {**state, "cv_parsed": cv_parsed, "status": "normalized"}


# ─── Node 4: Persist to DB ──────────────────────────────────────────────


async def persist_cv_data_node(state: CVParsingState) -> CVParsingState:
    """
    STEP 4: Persist parsed CV data to DB.
    - Update user_cvs: cv_parsed_json, full_name, summary, experience_years_total, status=completed
    - Upsert skills → skills + user_skill_profile tables
    """
    cv_id_str = state["cv_id"]
    cv_parsed: CVParsedData = state.get("cv_parsed")
    db = state["db"]

    logger.info(
        f"\n{'─' * 50}\n"
        f"[STEP 4] persist_cv_data_node | cv_id={cv_id_str}\n"
        f"  cv_parsed available: {bool(cv_parsed)}"
    )
    _update_cv_progress(cv_id_str, "Đang lưu trữ dữ liệu vào cơ sở dữ liệu...", 90)

    if not cv_parsed:
        logger.error(f"[STEP 4] ✗ FAIL — cv_parsed is None")
        return {**state, "status": "failed", "error": "No CV parsed data to persist"}

    # ── Validate cv_id ───────────────────────────────────────────────────────
    try:
        cv_uuid = uuid.UUID(cv_id_str)
    except ValueError as e:
        logger.error(f"[STEP 4] ✗ FAIL — invalid cv_id: {e}")
        return {**state, "status": "failed", "error": f"Invalid cv_id: {cv_id_str}"}

    # ── Load CV record ───────────────────────────────────────────────────────
    try:
        cv_record = db.query(UserCV).filter(UserCV.id == cv_uuid).first()
        if not cv_record:
            logger.error(f"[STEP 4] ✗ FAIL — CV record not found in DB for {cv_id_str}")
            return {**state, "status": "failed", "error": "CV record not found in DB"}
        logger.info(
            f"[STEP 4] CV record loaded | current status={cv_record.status} | "
            f"user_id={cv_record.user_id}"
        )
    except Exception as e:
        logger.error(f"[STEP 4] DB query failed: {e}", exc_info=True)
        return {**state, "status": "failed", "error": f"DB query failed: {e}"}

    from datetime import datetime

    try:
        # ── Update CV fields ──────────────────────────────────────────────────
        logger.info(
            f"[STEP 4] Updating UserCV fields:\n"
            f"  full_name           : {cv_parsed.get('full_name', '')[:50]}\n"
            f"  summary             : {(cv_parsed.get('summary') or '')[:80]}...\n"
            f"  experience_yrs      : {cv_parsed.get('experience_years_total')}\n"
            f"  seniority           : {cv_parsed.get('seniority')}\n"
            f"  skills count        : {len(cv_parsed.get('skills', []))}\n"
            f"  work_history count  : {len(cv_parsed.get('work_history', []))}\n"
            f"  education count     : {len(cv_parsed.get('education', []))}\n"
            f"  certifications      : {len(cv_parsed.get('certifications', []))}\n"
            f"  is_ocr              : {cv_parsed.get('is_ocr')}\n"
            f"  ocr_confidence      : {cv_parsed.get('ocr_confidence', 0):.2f}"
        )

        cv_record.cv_parsed_json = cv_parsed
        cv_record.cv_parsed_at = datetime.now()
        cv_record.experience_years_total = cv_parsed.get("experience_years_total", 0)
        cv_record.summary = cv_parsed.get("summary", "")
        cv_record.full_name = cv_parsed.get("full_name", "")
        cv_record.status = "completed"

        # Persist raw_text back to DB for cache hit on future parses
        if state.get("raw_text"):
            cv_record.raw_text = state["raw_text"]
            logger.info(
                f"[STEP 4] Cached raw_text ({len(state['raw_text'])} chars) in DB"
            )

        # ── Upsert Skills ────────────────────────────────────────────────────
        skills_upserted = await _upsert_skills_from_cv(
            cv_parsed.get("skills", []), cv_id_str, db
        )
        
        # ── Upsert Work History ──────────────────────────────────────────
        work_upserted = await _upsert_work_history_from_cv(
            cv_parsed.get("work_history", []), cv_id_str, db
        )

        # ── Commit ───────────────────────────────────────────────────────────
        db.commit()
        logger.info(
            f"[STEP 4] ✓ DB COMMIT successful | cv_id={cv_id_str}\n"
            f"  skills upserted : {skills_upserted}\n"
            f"  status          : 'completed'\n"
            f"{'─' * 50}"
        )

        # ── Send Email Notification ──────────────────────────────────────────
        try:
            from shared.models import User
            from shared.email_utils import notify_cv_parsing_complete
            
            user = db.query(User).filter(User.id == cv_record.user_id).first()
            if user and user.email:
                notify_cv_parsing_complete(
                    user_email=user.email,
                    cv_name=cv_record.full_name or "CV của bạn",
                    experience_years=cv_record.experience_years_total,
                    skills_count=len(cv_parsed.get('skills', []))
                )
                logger.info(f"[STEP 4] ✓ Email notification sent to {user.email}")
        except Exception as email_err:
            logger.warning(f"[STEP 4] Failed to send email notification: {email_err}")

        return {**state, "status": "persisted"}

    except Exception as e:
        db.rollback()
        logger.error(
            f"[STEP 4] ✗ FAIL — commit failed, rolled back: {e}", exc_info=True
        )
        return {**state, "status": "failed", "error": str(e)}


# ─── Helper: Upsert Skills ───────────────────────────────────────────────


async def _upsert_skills_from_cv(skills: list, cv_id: str, db) -> int:
    """
    Upsert skills từ parsed CV vào skills + user_skill_profile tables.
    Trả về số skills đã upsert.
    """
    from shared.models import Skill, UserSkillProfile

    upserted_count = 0
    skipped_count = 0
    error_count = 0
    cv_uuid = uuid.UUID(cv_id)

    logger.info(
        f"[STEP 4/SKILLS] Starting skill upsert | cv_id={cv_id} | "
        f"incoming skills={len(skills)}"
    )

    for idx, s in enumerate(skills):
        skill_name = (s.get("name") or "").strip()
        if not skill_name:
            skipped_count += 1
            logger.debug(f"  skill[{idx}] SKIPPED — empty name")
            continue

        # ── Upsert Skill record ──────────────────────────────────────────
        skill_record = db.query(Skill).filter(Skill.name == skill_name).first()

        if not skill_record:
            # Normalize category to English for DB consistency
            raw_cat = (s.get("category") or "Technology").strip()
            if raw_cat.lower() == "công nghệ":
                raw_cat = "Technology"

            skill_record = Skill(
                id=uuid.uuid4(), name=skill_name, category=raw_cat
            )
            db.add(skill_record)
            db.flush()
            logger.debug(
                f"  skill[{idx}] CREATED new Skill: id={skill_record.id} name='{skill_name}' category='{raw_cat}'"
            )
        else:
            logger.debug(
                f"  skill[{idx}] REUSED existing Skill: id={skill_record.id} name='{skill_name}'"
            )

        # ── Upsert UserSkillProfile ───────────────────────────────────────
        existing_profile = (
            db.query(UserSkillProfile)
            .filter(
                UserSkillProfile.cv_id == cv_uuid,
                UserSkillProfile.skill_id == skill_record.id,
            )
            .first()
        )

        profile_data = {
            "years_exp": max(0.0, float(s.get("experience_years") or s.get("years_exp") or 0)),
            "level": s.get("level", "Mid-level"),
            "source": "cv_parsed_v3",
        }

        if existing_profile:
            for key, val in profile_data.items():
                setattr(existing_profile, key, val)
            logger.debug(
                f"  skill[{idx}] UPDATED UserSkillProfile: "
                f"years_exp={profile_data['years_exp']} level={profile_data['level']}"
            )
        else:
            new_profile = UserSkillProfile(
                id=uuid.uuid4(),
                skill_id=skill_record.id,
                cv_id=cv_uuid,
                **profile_data,
            )
            db.add(new_profile)
            logger.debug(
                f"  skill[{idx}] INSERTED UserSkillProfile: "
                f"years_exp={profile_data['years_exp']} level={profile_data['level']}"
            )

        upserted_count += 1
        logger.debug(f"  skill[{idx}] ✓ upserted: name='{skill_name}'")

    logger.info(
        f"[STEP 4/SKILLS] Skill upsert complete | cv_id={cv_id}\n"
        f"  total incoming : {len(skills)}\n"
        f"  upserted       : {upserted_count}\n"
        f"  skipped (empty): {skipped_count}"
    )

    return upserted_count


async def _upsert_work_history_from_cv(work_history: list, cv_id: str, db) -> int:
    """
    BUG-005 FIX: Upsert work history with PII masking before database storage.
    """
    from shared.models import UserWorkExperience
    from shared.work_history_masking import mask_work_history
    
    upserted_count = 0
    cv_uuid = uuid.UUID(cv_id)
    
    # Xóa cũ ghi mới (đồng bộ với finalize_cv)
    db.query(UserWorkExperience).filter(UserWorkExperience.cv_id == cv_uuid).delete()
    
    # BUG-005 FIX: Apply PII masking to work history
    masked_work_history = mask_work_history(work_history or [])
    
    for w in masked_work_history:
        new_work = UserWorkExperience(
            id=uuid.uuid4(),
            cv_id=cv_uuid,
            position_name=w.get("position", "N/A"),
            company_name=w.get("company", "N/A"),
            duration_years=float(w.get("duration_years") or 0),
            description=w.get("description", "")
        )
        db.add(new_work)
        upserted_count += 1
        
    return upserted_count
